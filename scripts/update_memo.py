#!/usr/bin/env python3
"""Update the TMG Vetting Pipeline Memo with Session 7 changes."""
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

SRC = "/Users/shannonwheatman/Downloads/Recents/TMG_Vetting_Pipeline_Memo.docx"
DST = "/Users/shannonwheatman/Downloads/TMG_Vetting_Pipeline_Memo.docx"

doc = Document(SRC)

# ── Helper: find paragraph index by partial text match ──
def find_para(text, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if text in p.text:
            return i
    return None

# ── Helper: insert paragraph after index ──
def insert_paragraph_after(idx, text, style='Normal'):
    """Insert a new paragraph after paragraph at idx. Returns the new paragraph."""
    ref = doc.paragraphs[idx]._element
    new_p = docx.oxml.OxmlElement('w:p')
    ref.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, doc)
    new_para.style = doc.styles[style]
    new_para.text = text
    return new_para

# ══════════════════════════════════════════════════════════
# 1. UPDATE EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════

# Fix "10 different government databases" → "12+"
idx = find_para("TMG has built an automated due diligence pipeline")
if idx is not None:
    p = doc.paragraphs[idx]
    p.clear()
    p.add_run(
        "TMG has built an automated due diligence pipeline — internally called TMG Sentinel — "
        "that conducts comprehensive background vetting on prospective clients (individuals or "
        "organizations). The pipeline searches 12 different government databases, court systems, "
        "financial disclosure registries, federal contract records, and news sources, then uses AI "
        "to synthesize the findings into a "
    )
    bold_run = p.add_run("Combined Decision")
    bold_run.bold = True
    p.add_run(
        " based on the more cautious of two independent scores:"
    )

# Fix "approximately 10-20 minutes" → update timing
idx2 = find_para("The pipeline runs in approximately")
if idx2 is not None:
    p = doc.paragraphs[idx2]
    p.clear()
    p.add_run(
        "The pipeline runs in approximately 15-25 minutes for a standard vetting and produces a "
        "detailed, sourced report with an executive summary, scored risk dimensions, red/yellow flags, "
        "and a Combined Decision recommendation (Approve, Conditional Approve, Further Review, or "
        "Recommend Reject). The Combined Decision always takes the more cautious of the two scores — "
        "so a subject who is clean on paper but reputationally toxic still gets flagged."
    )

# ══════════════════════════════════════════════════════════
# 2. UPDATE STEP 4 — add pagination note
# ══════════════════════════════════════════════════════════
idx = find_para("Queries CourtListener RECAP — the largest open archive")
if idx is not None:
    p = doc.paragraphs[idx]
    p.clear()
    p.add_run(
        "Queries CourtListener RECAP — the largest open archive of federal court records (sourced "
        "from PACER). Searches both dockets (case records) and judicial opinions. Paginates through "
        "all results (up to 500 per search, safety-capped at 25 pages). Automatically classifies "
        "each case by risk level:"
    )

# ══════════════════════════════════════════════════════════
# 3. UPDATE STEP 9 — add pagination note
# ══════════════════════════════════════════════════════════
idx = find_para("Queries CourtListener with a filter for all federal bankruptcy courts")
if idx is not None:
    p = doc.paragraphs[idx]
    p.clear()
    p.add_run(
        "Queries CourtListener with a filter for all federal bankruptcy courts. Paginates through "
        "all results (up to 500, safety-capped at 25 pages). Classifies cases by chapter type "
        "(Chapter 7 liquidation, Chapter 11 reorganization, Chapter 13 individual repayment, "
        "Chapter 15 cross-border). Identifies whether the subject was the debtor (filed for "
        "bankruptcy) or a creditor (was owed money)."
    )

# ══════════════════════════════════════════════════════════
# 4. UPDATE STEP 10 — add org-specific queries
# ══════════════════════════════════════════════════════════
idx = find_para("Runs 15 targeted Tavily queries with advanced search depth")
if idx is not None:
    p = doc.paragraphs[idx]
    p.clear()
    p.add_run(
        "Runs 15 targeted Tavily queries with advanced search depth covering: controversy/scandal/"
        "investigation, lawsuits/fraud, criminal indictment/conviction, sanctions violations, "
        "political donations/lobbying, SEC enforcement, bankruptcy/financial problems, business "
        "dealings, social media statements, international connections, corruption/bribery, company "
        "reputation, and Twitter/X activity."
    )
    # Add org-specific note after the second paragraph of Step 10
    idx2 = find_para("Uses two search passes", idx)
    if idx2 is not None:
        p2 = doc.paragraphs[idx2]
        p2.clear()
        p2.add_run(
            "Uses two search passes: one using investigative source domains (ProPublica, ICIJ, OCCRP, "
            "quality newspapers), and a second open search with global excludes. Typically returns "
            "80-150 unique sources for well-known subjects. For organizations, 6 additional targeted "
            "queries are added covering: boycott/protest activity, CEO/executive public statements, "
            "government contract controversies, political donation patterns, employee dissent/"
            "whistleblower reports, and surveillance/enforcement/civil liberties concerns."
        )

# ══════════════════════════════════════════════════════════
# 5. ADD STEP 11 — Executive Identification (after Step 10)
# ══════════════════════════════════════════════════════════
# Find Step 12 heading and insert Step 11 before it
idx = find_para("Step 12 — International Checks")
if idx is not None:
    # We need to insert before Step 12. Insert using XML manipulation.
    ref_element = doc.paragraphs[idx]._element

    # Create Step 11 heading
    h3 = docx.oxml.OxmlElement('w:p')
    ref_element.addprevious(h3)
    h3_para = docx.text.paragraph.Paragraph(h3, doc)
    h3_para.style = doc.styles['Heading 3']
    h3_para.text = "Step 11 — Executive Identification & Mini-Vet (Organizations only)"

    # Create description paragraph
    desc = docx.oxml.OxmlElement('w:p')
    h3.addnext(desc)
    desc_para = docx.text.paragraph.Paragraph(desc, doc)
    desc_para.style = doc.styles['Normal']
    desc_para.text = (
        "For organization subjects, identifies top executives through SEC EDGAR Form 3 and Form 4 "
        "(insider ownership) filings. Looks up the company's SEC CIK number, retrieves recent "
        "ownership filings, then parses the XML to extract officer names, titles, and director "
        "status. For private companies not in SEC, falls back to Tavily web search to identify "
        "leadership."
    )

    # Create mini-vet paragraph
    mini = docx.oxml.OxmlElement('w:p')
    desc.addnext(mini)
    mini_para = docx.text.paragraph.Paragraph(mini, doc)
    mini_para.style = doc.styles['Normal']
    mini_para.text = (
        "Each identified executive then receives a mini due-diligence check: FEC campaign "
        "contribution search (who they donate to, how much), targeted Tavily news search "
        "(10 headlines per executive), and OpenSanctions screening. This surfaces individual "
        "political donation patterns and personal controversies that may not appear in the "
        "company-level search. For example, the Palantir vetting identified 8 executives "
        "and discovered CEO Alex Karp had $275,400 in FEC contributions heavily skewed "
        "Republican (including $100K to Cornyn Victory Committee)."
    )

    # Create limitation paragraph
    lim = docx.oxml.OxmlElement('w:p')
    mini.addnext(lim)
    lim_para = docx.text.paragraph.Paragraph(lim, doc)
    lim_para.style = doc.styles['Normal']
    lim_para.text = (
        "Limitation: Automatically skipped for individual subjects (executives are only relevant "
        "when vetting an organization). SEC EDGAR coverage is limited to public companies and "
        "their insiders — private company executive identification relies on web search."
    )

# ══════════════════════════════════════════════════════════
# 6. ADD STEP 14 — Government Contracts (after Step 12)
# ══════════════════════════════════════════════════════════
# Find "Phase 3: AI Synthesis" heading and insert Step 14 before it
idx = find_para("Phase 3: AI Synthesis")
if idx is not None:
    ref_element = doc.paragraphs[idx]._element

    h3 = docx.oxml.OxmlElement('w:p')
    ref_element.addprevious(h3)
    h3_para = docx.text.paragraph.Paragraph(h3, doc)
    h3_para.style = doc.styles['Heading 3']
    h3_para.text = "Step 14 — Federal Government Contracts"

    desc = docx.oxml.OxmlElement('w:p')
    h3.addnext(desc)
    desc_para = docx.text.paragraph.Paragraph(desc, doc)
    desc_para.style = doc.styles['Normal']
    desc_para.text = (
        "Searches USAspending.gov — the U.S. government's official database of all federal "
        "spending — for contract awards matching the subject's company name. Covers both prime "
        "contracts and indefinite delivery vehicles (IDVs). Paginates through up to 1,000 results. "
        "Aggregates contract data by awarding agency, providing total dollar amounts, award counts, "
        "and descriptions of the largest contracts. No API key required."
    )

    desc2 = docx.oxml.OxmlElement('w:p')
    desc.addnext(desc2)
    desc2_para = docx.text.paragraph.Paragraph(desc2, doc)
    desc2_para.style = doc.styles['Normal']
    desc2_para.text = (
        "This data is valuable for two reasons: it reveals the scale and nature of a subject's "
        "government relationships (e.g., $145M ICE contract, $2.3B DoD contracts), and it feeds "
        "into both the corporate/business risk dimension and the Reputational Contagion Score. "
        "Contracts with politically sensitive agencies (ICE, CBP, surveillance programs) are "
        "particularly relevant for TMG's reputational analysis."
    )

# ══════════════════════════════════════════════════════════
# 7. UPDATE SYNTHESIS SECTION — mention combined decision
# ══════════════════════════════════════════════════════════
idx = find_para("The pipeline sends all collected data from all previous steps")
if idx is not None:
    p = doc.paragraphs[idx]
    p.clear()
    p.add_run(
        "The pipeline sends all collected data from all previous steps — including executive "
        "donation profiles and government contract records — to Anthropic's Claude AI (specifically "
        "Claude Opus, the most capable model available) with a detailed prompt that instructs it to:"
    )

# Update the numbered list items after synthesis intro
idx = find_para("Score 7 risk dimensions")
if idx is not None:
    doc.paragraphs[idx].text = "1. Score 7 risk dimensions (0-10 each) based on the evidence"

idx = find_para("Score 6 reputational contagion questions")
if idx is not None:
    doc.paragraphs[idx].text = "2. Score 6 reputational contagion questions (0-10 each) based on TMG's specific identity"

idx = find_para("Make a recommendation")
if idx is not None:
    doc.paragraphs[idx].text = "6. Make a recommendation (the Combined Decision — the more cautious of factual risk vs. reputational risk)"

# ══════════════════════════════════════════════════════════
# 8. ADD COMBINED DECISION SECTION (after Divergence Alerts)
# ══════════════════════════════════════════════════════════
idx = find_para("Divergence Alerts:")
if idx is not None:
    ref = doc.paragraphs[idx]._element

    h2 = docx.oxml.OxmlElement('w:p')
    ref.addnext(h2)
    h2_para = docx.text.paragraph.Paragraph(h2, doc)
    h2_para.style = doc.styles['Heading 2']
    h2_para.text = "Score 3: The Combined Decision"

    desc = docx.oxml.OxmlElement('w:p')
    h2.addnext(desc)
    desc_para = docx.text.paragraph.Paragraph(desc, doc)
    desc_para.style = doc.styles['Normal']
    desc_para.text = (
        "The Combined Decision is the pipeline's final, authoritative recommendation. It automatically "
        "takes the MORE CAUTIOUS of the Factual Risk Score and the Reputational Contagion Score. "
        "This means a subject cannot \"sneak through\" on clean legal history if their reputational "
        "profile is toxic for TMG, and vice versa."
    )

    desc2 = docx.oxml.OxmlElement('w:p')
    desc.addnext(desc2)
    desc2_para = docx.text.paragraph.Paragraph(desc2, doc)
    desc2_para.style = doc.styles['Normal']
    desc2_para.text = (
        "The Combined Decision includes: the recommendation (Approve through Recommend Reject), "
        "the combined tier (LOW through CRITICAL), which score drove the decision (factual_risk "
        "or reputational_contagion), and an explanation of why. When the two scores diverge "
        "significantly — e.g., factual risk is MODERATE but RCS is HIGH — a Divergence Alert is "
        "triggered, explicitly warning that traditional due diligence would have missed the risk."
    )

    desc3 = docx.oxml.OxmlElement('w:p')
    desc2.addnext(desc3)
    desc3_para = docx.text.paragraph.Paragraph(desc3, doc)
    desc3_para.style = doc.styles['Normal']
    desc3_para.text = (
        "Example: Palantir Technologies scored 2.74/10 MODERATE on factual risk (Conditional Approve) "
        "but 7.85/10 HIGH on the RCS. The Combined Decision was HIGH — Recommend Reject, requiring "
        "unanimous partner approval to override. The driver was reputational contagion. Without the "
        "dual-score system, traditional due diligence would have recommended approval."
    )

# ══════════════════════════════════════════════════════════
# 9. UPDATE REAL EXAMPLES
# ══════════════════════════════════════════════════════════

# --- Peter Thiel ---
idx = find_para("Peter Thiel (Individual, Domestic Political)")
if idx is not None:
    # Update the bullet points after the heading
    idx2 = find_para("Quick Screen Result:", idx)
    if idx2 is not None:
        doc.paragraphs[idx2].text = (
            "Deep Dive Result: Factual Risk 5.50/10 — ELEVATED (Further Review) | "
            "RCS 8.72/10 — CRITICAL | Combined Decision: CRITICAL — Recommend Reject"
        )

    idx3 = find_para("Sanctions gate: PASS", idx)
    if idx3 is not None:
        doc.paragraphs[idx3].text = (
            "Sanctions gate: PASS. Debarment gate: PASS."
        )

    idx4 = find_para("114 Tavily news sources", idx)
    if idx4 is not None:
        doc.paragraphs[idx4].text = (
            "114 Tavily news sources, 257 court dockets, 387 EDGAR filings, 533 FEC contributions, "
            "1,950 SEC filings, bankruptcy records, international checks collected"
        )

    # Update the summary paragraph
    idx5 = find_para("The factual score was moderate", idx)
    if idx5 is not None:
        p = doc.paragraphs[idx5]
        p.clear()
        p.add_run(
            "The factual score was ELEVATED (5.50/10) — no criminal cases, but significant litigation "
            "risk, media controversies (IRS investigation into $5B Roth IRA tax avoidance, congressional "
            "investigation into CCP propaganda network), and high political/lobbying exposure. But the "
            "real story is the RCS: 8.72/10 CRITICAL. Every RCS question scored 7.5 or above — "
            "Q1 Partisan Alignment at 9.5 (major MAGA donor, JD Vance backer), Q3 Narrative Vulnerability "
            "at 9.0 (\"Obama's Campaign Manager Now Working for Trump Megadonor Behind Mass Surveillance\"), "
            "and Q5 Industry Toxicity at 9.0 (surveillance tech used against civil liberties). The Combined "
            "Decision is CRITICAL — Recommend Reject. This is the dual-score system working exactly as designed."
        )

# --- Palantir ---
idx = find_para("Palantir Technologies (Organization, Domestic Corporate)")
if idx is not None:
    # Update result line
    idx2 = find_para("Standard Vet Result:", idx)
    if idx2 is not None:
        p = doc.paragraphs[idx2]
        p.clear()
        p.add_run(
            "Deep Dive Result: Factual Risk 2.74/10 — MODERATE (Conditional Approve) | "
            "RCS 7.85/10 — HIGH | Combined Decision: HIGH — Recommend Reject (requires "
            "unanimous partner approval)"
        )

    idx3 = find_para("98 Tavily sources", idx)
    if idx3 is not None:
        doc.paragraphs[idx3].text = (
            "148 Tavily sources, 248 court dockets, 3,166 EDGAR filings, 10,000+ SEC filings, "
            "8 executives identified and mini-vetted, 414 government contracts ($3.94B total) collected"
        )

    idx4 = find_para("Factual risk was low", idx)
    if idx4 is not None:
        doc.paragraphs[idx4].text = (
            "Factual risk was moderate — the company itself has no sanctions, no debarments, "
            "but has an active securities fraud class action (Oct 2025), trade secret litigation, "
            "and 248 court dockets. A Divergence Alert was triggered."
        )

    # Update the RCS bullet points
    idx5 = find_para("Partisan Alignment: 3/10", idx)
    if idx5 is not None:
        doc.paragraphs[idx5].text = (
            "Partisan Alignment: 7.5/10 (Thiel is a prominent Republican donor; Karp's FEC records "
            "show $275K heavily skewed GOP including $100K to Cornyn Victory Committee)"
        )

    idx6 = find_para("Stakeholder Backlash: 6/10", idx)
    if idx6 is not None:
        doc.paragraphs[idx6].text = (
            "Stakeholder Backlash: 8.0/10 (DCCC/DNC would view engagement as betrayal; former "
            "employees published open letter condemning Trump admin immigration work)"
        )

    idx7 = find_para("Narrative Vulnerability: 7/10", idx)
    if idx7 is not None:
        doc.paragraphs[idx7].text = (
            "Narrative Vulnerability: 9.0/10 (\"Obama's Former Campaign Manager Now Helping "
            "Trump-Backed Surveillance Company Target Immigrants\")"
        )

    idx8 = find_para("Industry Toxicity: 7/10", idx)
    if idx8 is not None:
        doc.paragraphs[idx8].text = (
            "Industry Toxicity: 8.5/10 (surveillance tech for immigration enforcement; CEO used "
            "slur to describe critics; company profits from authoritarian applications)"
        )

    idx9 = find_para("Temporal Context: 5/10", idx)
    if idx9 is not None:
        doc.paragraphs[idx9].text = (
            "Temporal Context: 7.0/10 (active congressional investigations, Pentagon AI disputes, "
            "Mandelson/Epstein connection exposed, NHS deal controversies in UK)"
        )

    # Update the conclusion paragraph
    idx10 = find_para("This is exactly the kind of case the dual-score system", idx)
    if idx10 is not None:
        p = doc.paragraphs[idx10]
        p.clear()
        p.add_run(
            "This is exactly the kind of case the dual-score system was designed for. Traditional "
            "due diligence says \"conditional approve\" (2.74/10 MODERATE). The RCS says \"recommend "
            "reject\" (7.85/10 HIGH). The Combined Decision takes the more cautious signal: HIGH — "
            "Recommend Reject, requiring unanimous partner approval to override. The pipeline also "
            "identified 8 key executives, revealing that Karp ($275K), Sankar ($66K), and Thiel "
            "($61K) donate almost exclusively to Republican candidates — data that was invisible "
            "in the company-level FEC search. Government contract analysis showed $3.94 billion "
            "across 15 agencies, including the politically sensitive $145M ICE contract."
        )

# ══════════════════════════════════════════════════════════
# 10. UPDATE VETTING LEVELS TABLE (Table 4)
# ══════════════════════════════════════════════════════════
table4 = doc.tables[4]
# Update Standard Vet row
table4.rows[2].cells[1].text = (
    "All of Quick Screen + Litigation, Corporate, FEC, SEC, Lobbying, "
    "Deep News (15+ queries), Executive ID & Mini-Vet, Government Contracts"
)
table4.rows[2].cells[2].text = "~15-20 minutes"

# Update Deep Dive row
table4.rows[3].cells[1].text = (
    "All of Standard Vet + Bankruptcy, International/PEP"
)
table4.rows[3].cells[2].text = "~20-30 minutes"

# ══════════════════════════════════════════════════════════
# 11. UPDATE DATA SOURCES TABLE (Table 5) — add USAspending.gov and EDGAR Form 3/4
# ══════════════════════════════════════════════════════════
table5 = doc.tables[5]

# Add USAspending.gov row
new_row = table5.add_row()
new_row.cells[0].text = "USAspending.gov"
new_row.cells[1].text = "Free (no key needed)"
new_row.cells[2].text = "Federal contract awards — agency, amount, description for all government contracts"

# Add SEC EDGAR Form 3/4 row (exec identification)
new_row2 = table5.add_row()
new_row2.cells[0].text = "SEC EDGAR (Form 3/4)"
new_row2.cells[1].text = "Free (public API)"
new_row2.cells[2].text = "Insider ownership filings — identifies company officers and directors for executive mini-vetting"

# ══════════════════════════════════════════════════════════
# 12. UPDATE COST TABLE (Table 6)
# ══════════════════════════════════════════════════════════
table6 = doc.tables[6]
# Update Tavily row to reflect more queries
table6.rows[2].cells[0].text = "Tavily searches (15-21 queries for standard vet)"
table6.rows[2].cells[1].text = "~$1-4 total"

# Update total
table6.rows[5].cells[1].text = "$3-8"

# Add exec mini-vet row before total
# Can't easily insert row before last, so update total text
table6.rows[5].cells[0].text = "Total per standard vetting (includes exec mini-vet Tavily + sanctions queries)"

# ══════════════════════════════════════════════════════════
# 13. UPDATE "CURRENTLY AUTOMATED" BULLET LIST
# ══════════════════════════════════════════════════════════
idx = find_para("AI-powered risk synthesis with dual scoring")
if idx is not None:
    p = doc.paragraphs[idx]
    p.text = "AI-powered risk synthesis with dual scoring and Combined Decision"
idx_nonprofit = find_para("Nonprofit/NGO connections")
if idx_nonprofit is not None:
    ref = doc.paragraphs[idx_nonprofit]._element

    b1 = docx.oxml.OxmlElement('w:p')
    ref.addnext(b1)
    b1_para = docx.text.paragraph.Paragraph(b1, doc)
    b1_para.style = doc.styles['List Bullet']
    b1_para.text = "Executive identification & mini-vetting (FEC donations, news, sanctions per exec)"

    b2 = docx.oxml.OxmlElement('w:p')
    b1.addnext(b2)
    b2_para = docx.text.paragraph.Paragraph(b2, doc)
    b2_para.style = doc.styles['List Bullet']
    b2_para.text = "Federal government contract records (USAspending.gov)"

# ══════════════════════════════════════════════════════════
# 14. UPDATE DATE
# ══════════════════════════════════════════════════════════
idx = find_para("Prepared by: Shannon Wheatman")
if idx is not None:
    doc.paragraphs[idx].text = (
        "Prepared by: Shannon Wheatman\n"
        "Date: March 8, 2026 (Updated)\n"
        "Distribution: TMG Leadership & Vetting Team\n"
        "Classification: Internal — Confidential"
    )

# ══════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════
doc.save(DST)
print(f"Saved updated memo to: {DST}")
